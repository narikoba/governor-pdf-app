# main.py
import streamlit as st
import pdfplumber
from io import BytesIO
from datetime import datetime
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openai import OpenAI
import tiktoken
import os

# OpenAI APIクライアント初期化（v1対応）
client = OpenAI(api_key=st.secrets["openai_api_key"])

st.title("知事記者会見録 自動整形アプリ（高精度版・有料会員向け）")

uploaded_file = st.file_uploader("PDFファイルをアップロードしてください", type="pdf")

if uploaded_file:
    # ファイル名から日付抽出（カッコ内も含めてOKに）
    match = re.search(r"[\(\[]?(\d{4})(\d{2})(\d{2})[\)\]]?", uploaded_file.name)
    if match:
        year, month, day = match.groups()
        formatted_date = f"{int(year)}.{int(month)}.{int(day)}"
        output_filename = f"{formatted_date}知事記者会見.docx"
    else:
        st.error("ファイル名に日付が含まれていません（例：20250328）。")
        st.stop()

    # PDFからテキストを抽出
    with pdfplumber.open(uploaded_file) as pdf:
        text = "\n".join(page.extract_text() for page in pdf.pages[1:] if page.extract_text())

    # トークン数で制限（gpt-4上限を考慮して4000トークン以内に）
    encoding = tiktoken.encoding_for_model("gpt-4")
    tokens = encoding.encode(text)
    max_tokens = 4000
    text = encoding.decode(tokens[:max_tokens])

    # ChatGPTに渡すプロンプト（精度向上版）
    prompt = f"""
以下の東京都知事会見録のテキストを、レイアウトと読みやすさを整えてください。

【目的】
都知事会見の記録として、見出しを明確にし、各話題が読みやすいように整えてください。読み手が内容をスムーズに理解できるよう工夫してください。

【編集方針】
- 話題ごとに「◉〇〇〇（タイトル）」という見出しを付けてください
- 話題の切れ目が分かるよう、段落を分けてください
- 発言者（知事、記者など）は【知事】のように記載し、行頭で改行してください
- 【知事】や【記者】は太字で表示する前提で構成してください
- インデント、行間を自然に整えてください
- 冗長な言い回しや繰り返し表現は削除・言い換えてください
- 以下のような定型文は削除してください：
    ・「詳細については〇〇へお問い合わせください」
    ・「会見で使用したスライド資料はこちらからご覧いただけます」
- 日本語として自然で、公式記録として適した文章にしてください

---
{text}
    """

    with st.spinner("ChatGPT（gpt-4）で文章を整形中..."):
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "あなたは日本語の行政文書編集の専門家です。"},
                {"role": "user", "content": prompt}
            ],
            temperature=0.4,
        )
        cleaned_text = response.choices[0].message.content

    # Wordファイルとして保存（見た目を整える）
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "MS Gothic"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "MS Gothic")
    style.font.size = Pt(11)

    for line in cleaned_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("◉"):
            doc.add_heading(line.strip("◉ "), level=2)
        elif line.startswith("【知事】") or line.startswith("【記者】") or line.startswith("【司会】"):
            p = doc.add_paragraph()
            run_bold = p.add_run(line[:4])  # 【知事】など
            run_bold.bold = True
            p.add_run(line[4:].lstrip())
        else:
            doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.success("整形が完了しました！")
    st.download_button("Wordファイルをダウンロード", buffer, file_name=output_filename)
